#!/usr/bin/env python3
"""
Test script for Office document support in Whisper Network.
Creates sample Office files and tests anonymization.

Tests:
- .docx (Microsoft Word)
- .xlsx (Microsoft Excel)
- .odt (LibreOffice Writer)
- .ods (LibreOffice Calc)

Author: Sylvain JOLY, NANO by NXO
"""

import asyncio
import requests
import os
from pathlib import Path
import io

# API Configuration
API_URL = "http://localhost:8001"
API_KEY = "dev_test_key_12345"


def create_test_docx(filename: str):
    """Create a test .docx file with sample content."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Rapport Confidentiel', 0)
        
        # Add paragraphs with sensitive data
        doc.add_paragraph(
            "Ce rapport concerne M. Jean Dupont, employé ID-12345, "
            "travaillant au département R&D."
        )
        
        doc.add_paragraph(
            "Mme Marie Martin (ID-67890) a signalé un incident réseau "
            "provenant de l'adresse IP 192.168.1.100."
        )
        
        doc.add_paragraph(
            "Contact: jean.dupont@example.com, tel: +33 6 12 34 56 78"
        )
        
        # Add a table
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Nom'
        header_cells[1].text = 'Email'
        header_cells[2].text = 'IP'
        
        # Data rows
        row1 = table.rows[1].cells
        row1[0].text = 'Alice Durand'
        row1[1].text = 'alice.durand@company.fr'
        row1[2].text = '10.0.0.25'
        
        row2 = table.rows[2].cells
        row2[0].text = 'Bob Martin'
        row2[1].text = 'bob.martin@company.fr'
        row2[2].text = '192.168.2.50'
        
        # Save
        doc.save(filename)
        print(f"✅ Created {filename}")
        return True
    except Exception as e:
        print(f"❌ Failed to create {filename}: {e}")
        return False


def create_test_xlsx(filename: str):
    """Create a test .xlsx file with sample content."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Employés"
        
        # Header
        headers = ['ID', 'Nom', 'Email', 'Téléphone', 'IP Bureau']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
        
        # Sample data
        data = [
            ['EMP-001', 'Sophie Dubois', 'sophie.dubois@corp.fr', '01 23 45 67 89', '192.168.1.10'],
            ['EMP-002', 'Thomas Bernard', 'thomas.bernard@corp.fr', '01 98 76 54 32', '192.168.1.11'],
            ['EMP-003', 'Claire Petit', 'claire.petit@corp.fr', '06 11 22 33 44', '10.0.0.15'],
            ['EMP-004', 'Marc Roux', 'marc.roux@corp.fr', '+33 6 55 44 33 22', '172.16.0.20'],
        ]
        
        for row_idx, row_data in enumerate(data, start=2):
            for col_idx, value in enumerate(row_data, start=1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        # Add a second sheet
        ws2 = wb.create_sheet("Serveurs")
        ws2['A1'] = 'Serveur'
        ws2['B1'] = 'IP'
        ws2['C1'] = 'Responsable'
        
        ws2['A2'] = 'SRV-WEB-01'
        ws2['B2'] = '192.168.100.10'
        ws2['C2'] = 'admin@company.com'
        
        ws2['A3'] = 'SRV-DB-01'
        ws2['B3'] = '10.20.30.40'
        ws2['C3'] = 'dba@company.com'
        
        wb.save(filename)
        print(f"✅ Created {filename}")
        return True
    except Exception as e:
        print(f"❌ Failed to create {filename}: {e}")
        return False


def create_test_odt(filename: str):
    """Create a test .odt file with sample content."""
    try:
        from odf.opendocument import OpenDocumentText
        from odf.text import P, H
        from odf.style import Style, TextProperties
        from odf import teletype
        
        doc = OpenDocumentText()
        
        # Title
        h = H(outlinelevel=1, text="Compte-rendu de réunion")
        doc.text.addElement(h)
        
        # Paragraphs with sensitive data
        p1 = P()
        teletype.addTextToElement(p1, "Présents: M. Pierre Lefevre, Mme Anne Moreau, M. Lucas Garnier")
        doc.text.addElement(p1)
        
        p2 = P()
        teletype.addTextToElement(p2, "Discussion sur l'incident réseau du 15/11/2025 concernant le serveur 192.168.50.100")
        doc.text.addElement(p2)
        
        p3 = P()
        teletype.addTextToElement(p3, "Contact IT: support@entreprise.fr ou tel +33 1 42 53 64 75")
        doc.text.addElement(p3)
        
        p4 = P()
        teletype.addTextToElement(p4, "Employé ID-99999 (Jacques Dumont) a été informé de la procédure.")
        doc.text.addElement(p4)
        
        doc.save(filename)
        print(f"✅ Created {filename}")
        return True
    except Exception as e:
        print(f"❌ Failed to create {filename}: {e}")
        return False


def create_test_ods(filename: str):
    """Create a test .ods file with sample content."""
    try:
        from odf.opendocument import OpenDocumentSpreadsheet
        from odf.table import Table, TableRow, TableCell
        from odf.text import P
        
        doc = OpenDocumentSpreadsheet()
        
        # Create table
        table = Table(name="Contacts")
        
        # Header row
        tr = TableRow()
        for header in ['Nom', 'Email', 'Téléphone', 'Département']:
            tc = TableCell()
            p = P(text=header)
            tc.addElement(p)
            tr.addElement(tc)
        table.addElement(tr)
        
        # Data rows
        contacts = [
            ['Julien Blanc', 'julien.blanc@societe.fr', '06 12 34 56 78', 'Marketing'],
            ['Emma Rousseau', 'emma.rousseau@societe.fr', '01 23 45 67 89', 'Finance'],
            ['Noah Lambert', 'noah.lambert@societe.fr', '+33 6 98 76 54 32', 'IT'],
        ]
        
        for contact in contacts:
            tr = TableRow()
            for value in contact:
                tc = TableCell()
                p = P(text=value)
                tc.addElement(p)
                tr.addElement(tc)
            table.addElement(tr)
        
        doc.spreadsheet.addElement(table)
        doc.save(filename)
        print(f"✅ Created {filename}")
        return True
    except Exception as e:
        print(f"❌ Failed to create {filename}: {e}")
        return False


async def test_file_anonymization(filename: str):
    """Test anonymization of a single file."""
    try:
        print(f"\n🔄 Testing {filename}...")
        
        # Read file
        with open(filename, 'rb') as f:
            file_content = f.read()
        
        print(f"   📂 File size: {len(file_content)} bytes")
        
        # Upload and anonymize
        files = {'file': (filename, file_content)}
        headers = {'X-API-Key': API_KEY}
        
        response = requests.post(
            f"{API_URL}/anonymize-file",
            files=files,
            headers=headers,
            timeout=30
        )
        
        if response.status_code == 200:
            # Success - metadata is in headers, file is in content
            anon_count = response.headers.get('X-Anonymizations-Count', '0')
            proc_time = response.headers.get('X-Processing-Time-Ms', '0')
            file_type = response.headers.get('X-File-Type', 'unknown')
            new_filename = response.headers.get('Content-Disposition', '').split('filename="')[-1].rstrip('"')
            
            print(f"   ✅ Anonymization successful!")
            print(f"   📊 Anonymizations: {anon_count}")
            print(f"   ⏱️  Processing time: {proc_time}ms")
            print(f"   📁 File type: {file_type}")
            print(f"   📝 New filename: {new_filename}")
            
            # Save anonymized file
            output_filename = new_filename or filename.replace('.', '.anonymized.')
            with open(output_filename, 'wb') as f:
                f.write(response.content)
            print(f"   💾 Saved to {output_filename} ({len(response.content)} bytes)")
            
            return True
        else:
            print(f"   ❌ API Error {response.status_code}: {response.text}")
            return False
            
    except Exception as e:
        print(f"   ❌ Test failed: {e}")
        return False


async def main():
    """Main test runner."""
    print("=" * 60)
    print("🧪 Office Documents Anonymization Test")
    print("=" * 60)
    
    # Create test files directory
    test_dir = Path("test-office-files")
    test_dir.mkdir(exist_ok=True)
    os.chdir(test_dir)
    
    # Create test files
    print("\n📝 Creating test files...")
    test_files = []
    
    if create_test_docx("test_document.docx"):
        test_files.append("test_document.docx")
    
    if create_test_xlsx("test_spreadsheet.xlsx"):
        test_files.append("test_spreadsheet.xlsx")
    
    if create_test_odt("test_writer.odt"):
        test_files.append("test_writer.odt")
    
    if create_test_ods("test_calc.ods"):
        test_files.append("test_calc.ods")
    
    if not test_files:
        print("❌ No test files created!")
        return
    
    print(f"\n✅ Created {len(test_files)} test files")
    
    # Test anonymization
    print("\n" + "=" * 60)
    print("🔒 Testing Anonymization")
    print("=" * 60)
    
    success_count = 0
    for filename in test_files:
        if await test_file_anonymization(filename):
            success_count += 1
    
    # Summary
    print("\n" + "=" * 60)
    print("📊 Test Summary")
    print("=" * 60)
    print(f"✅ Successful: {success_count}/{len(test_files)}")
    print(f"❌ Failed: {len(test_files) - success_count}/{len(test_files)}")
    
    if success_count == len(test_files):
        print("\n🎉 All tests passed!")
    else:
        print("\n⚠️  Some tests failed")
    
    print(f"\n📁 Test files location: {test_dir.absolute()}")


if __name__ == "__main__":
    asyncio.run(main())
