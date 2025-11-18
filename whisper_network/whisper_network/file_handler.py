#!/usr/bin/env python3
"""
File Handler Module for Whisper Network.
Handles file upload, parsing, anonymization and export with format preservation.

Supports:
- Text files: .txt, .md, .log
- Config files: .conf, .ini, .yaml, .json, .toml, .env
- Scripts: .sh, .bash, .zsh, .ps1, .py, .js, .ts, .java, .cpp, .cs, etc.
- Office docs: .docx, .xlsx, .odt, .ods
- Rich formats: .pdf, .rtf

Developed by Sylvain JOLY, NANO by NXO
License: MIT
"""

import os
import logging
import mimetypes
import io
from typing import Optional, Dict, Any, Tuple
from pathlib import Path
import chardet
from dataclasses import dataclass
from enum import Enum

logger = logging.getLogger(__name__)


class FileType(Enum):
    """Supported file types."""
    TEXT = "text"
    CONFIG = "config"
    SCRIPT = "script"
    OFFICE = "office"
    RICH = "rich"
    UNKNOWN = "unknown"


@dataclass
class FileInfo:
    """Information about uploaded file."""
    filename: str
    size_bytes: int
    mime_type: str
    file_type: FileType
    extension: str
    encoding: str
    content: str


class FileHandler:
    """Handle file upload, parsing and export with format preservation."""
    
    # Maximum file size: 10MB
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # Supported extensions by category
    TEXT_EXTENSIONS = {'.txt', '.md', '.log', '.rst', '.csv'}
    CONFIG_EXTENSIONS = {'.conf', '.ini', '.yaml', '.yml', '.json', '.toml', '.env', '.properties', '.cfg'}
    SCRIPT_EXTENSIONS = {
        '.sh', '.bash', '.zsh', '.fish',  # Shell
        '.ps1', '.psm1', '.psd1',  # PowerShell
        '.py', '.pyw',  # Python
        '.js', '.mjs', '.cjs', '.ts', '.tsx',  # JavaScript/TypeScript
        '.java', '.kt', '.scala',  # JVM
        '.cpp', '.cc', '.cxx', '.c', '.h', '.hpp',  # C/C++
        '.cs', '.vb',  # .NET
        '.go', '.rs', '.swift', '.rb', '.php', '.pl', '.lua'  # Other languages
    }
    OFFICE_EXTENSIONS = {
        '.docx', '.doc',  # Microsoft Word
        '.xlsx', '.xls',  # Microsoft Excel
        '.odt',  # LibreOffice Writer
        '.ods',  # LibreOffice Calc
    }
    RICH_EXTENSIONS = {
        '.pdf',  # Portable Document Format
        '.rtf',  # Rich Text Format
    }
    
    # Supported MIME types
    ALLOWED_MIME_TYPES = {
        'text/plain',
        'text/markdown',
        'text/x-log',
        'text/x-python',
        'text/x-script',
        'text/x-shellscript',
        'application/x-sh',
        'application/x-bash',
        'application/json',
        'application/yaml',
        'application/x-yaml',
        'application/toml',
        'text/x-ini',
        'text/x-properties',
        # Office documents
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # .xlsx
        'application/vnd.oasis.opendocument.text',  # .odt
        'application/vnd.oasis.opendocument.spreadsheet',  # .ods
        'application/msword',  # .doc (old format - limited support)
        'application/vnd.ms-excel',  # .xls (old format - limited support)
        # Rich formats
        'application/pdf',  # .pdf
        'text/rtf',  # .rtf
        'application/rtf',  # .rtf (alternative)
    }
    
    def __init__(self):
        """Initialize file handler."""
        logger.info("FileHandler initialized")
    
    def _detect_encoding(self, file_bytes: bytes) -> str:
        """
        Detect file encoding using chardet.
        
        Args:
            file_bytes: Raw file bytes
            
        Returns:
            Detected encoding (defaults to utf-8)
        """
        try:
            result = chardet.detect(file_bytes)
            encoding = result.get('encoding', 'utf-8')
            confidence = result.get('confidence', 0)
            
            logger.debug(f"Detected encoding: {encoding} (confidence: {confidence:.2%})")
            
            # Fallback to utf-8 if confidence is too low
            if confidence < 0.5:
                logger.warning(f"Low encoding confidence ({confidence:.2%}), defaulting to utf-8")
                return 'utf-8'
            
            return encoding or 'utf-8'
        
        except Exception as e:
            logger.warning(f"Encoding detection failed: {e}, defaulting to utf-8")
            return 'utf-8'
    
    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        """
        Extract text from .docx file (Microsoft Word).
        
        Args:
            file_bytes: Raw .docx file bytes
            
        Returns:
            Extracted text content
        """
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        tables_text.append(row_text)
            
            # Combine all text
            all_text = '\n'.join(paragraphs)
            if tables_text:
                all_text += '\n\n' + '\n'.join(tables_text)
            
            logger.info(f"Extracted {len(paragraphs)} paragraphs from .docx")
            return all_text
        except Exception as e:
            logger.error(f"Failed to extract text from .docx: {e}")
            raise ValueError(f"Cannot parse .docx file: {str(e)}")
    
    def _extract_text_from_xlsx(self, file_bytes: bytes) -> str:
        """
        Extract text from .xlsx file (Microsoft Excel).
        
        Args:
            file_bytes: Raw .xlsx file bytes
            
        Returns:
            Extracted text content
        """
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            
            all_text = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                all_text.append(f"=== {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None and empty cells
                    row_values = [str(cell) for cell in row if cell is not None and str(cell).strip()]
                    if row_values:
                        all_text.append('\t'.join(row_values))
            
            result = '\n'.join(all_text)
            logger.info(f"Extracted {len(wb.sheetnames)} sheets from .xlsx")
            return result
        except Exception as e:
            logger.error(f"Failed to extract text from .xlsx: {e}")
            raise ValueError(f"Cannot parse .xlsx file: {str(e)}")
    
    def _extract_text_from_odt(self, file_bytes: bytes) -> str:
        """
        Extract text from .odt file (LibreOffice Writer).
        
        Args:
            file_bytes: Raw .odt file bytes
            
        Returns:
            Extracted text content
        """
        try:
            from odf import text, teletype
            from odf.opendocument import load
            
            doc = load(io.BytesIO(file_bytes))
            
            # Extract all text elements
            all_paragraphs = doc.getElementsByType(text.P)
            all_text = [teletype.extractText(para) for para in all_paragraphs if teletype.extractText(para).strip()]
            
            result = '\n'.join(all_text)
            logger.info(f"Extracted {len(all_text)} paragraphs from .odt")
            return result
        except Exception as e:
            logger.error(f"Failed to extract text from .odt: {e}")
            raise ValueError(f"Cannot parse .odt file: {str(e)}")
    
    def _extract_text_from_ods(self, file_bytes: bytes) -> str:
        """
        Extract text from .ods file (LibreOffice Calc).
        
        Args:
            file_bytes: Raw .ods file bytes
            
        Returns:
            Extracted text content
        """
        try:
            from odf import table, text, teletype
            from odf.opendocument import load
            
            doc = load(io.BytesIO(file_bytes))
            
            all_text = []
            tables = doc.spreadsheet.getElementsByType(table.Table)
            
            for tbl in tables:
                # Get table name
                table_name = tbl.getAttribute('name')
                if table_name:
                    all_text.append(f"=== {table_name} ===")
                
                # Extract rows
                rows = tbl.getElementsByType(table.TableRow)
                for row in rows:
                    cells = row.getElementsByType(table.TableCell)
                    row_values = []
                    for cell in cells:
                        cell_text = teletype.extractText(cell).strip()
                        if cell_text:
                            row_values.append(cell_text)
                    if row_values:
                        all_text.append('\t'.join(row_values))
            
            result = '\n'.join(all_text)
            logger.info(f"Extracted {len(tables)} tables from .ods")
            return result
        except Exception as e:
            logger.error(f"Failed to extract text from .ods: {e}")
            raise ValueError(f"Cannot parse .ods file: {str(e)}")
    
    def _extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """
        Extract text from .pdf file.
        Uses pdfplumber for better table/layout extraction, falls back to pypdf.
        
        Args:
            file_bytes: Raw .pdf file bytes
            
        Returns:
            Extracted text content
        """
        try:
            import pdfplumber
            from pypdf import PdfReader
            
            # Try pdfplumber first (better for tables and layout)
            try:
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    all_text = []
                    
                    for page_num, page in enumerate(pdf.pages, start=1):
                        # Extract text
                        text = page.extract_text()
                        if text and text.strip():
                            all_text.append(f"=== Page {page_num} ===")
                            all_text.append(text.strip())
                        
                        # Extract tables
                        tables = page.extract_tables()
                        if tables:
                            for table_idx, table in enumerate(tables, start=1):
                                all_text.append(f"\n--- Table {table_idx} ---")
                                for row in table:
                                    if row:
                                        row_text = '\t'.join([str(cell) if cell else '' for cell in row])
                                        if row_text.strip():
                                            all_text.append(row_text)
                    
                    result = '\n'.join(all_text)
                    if result.strip():
                        logger.info(f"Extracted text from PDF using pdfplumber ({len(pdf.pages)} pages)")
                        return result
            except Exception as plumber_error:
                logger.warning(f"pdfplumber extraction failed: {plumber_error}, trying pypdf")
            
            # Fallback to pypdf (simpler but works for most PDFs)
            reader = PdfReader(io.BytesIO(file_bytes))
            all_text = []
            
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    all_text.append(f"=== Page {page_num} ===")
                    all_text.append(text.strip())
            
            result = '\n'.join(all_text)
            if not result.strip():
                raise ValueError("PDF appears to be empty or contains only images (OCR not supported)")
            
            logger.info(f"Extracted text from PDF using pypdf ({len(reader.pages)} pages)")
            return result
        
        except Exception as e:
            logger.error(f"Failed to extract text from .pdf: {e}")
            raise ValueError(f"Cannot parse .pdf file: {str(e)}")
    
    def _extract_text_from_rtf(self, file_bytes: bytes) -> str:
        """
        Extract text from .rtf file (Rich Text Format).
        
        Args:
            file_bytes: Raw .rtf file bytes
            
        Returns:
            Extracted text content
        """
        try:
            from striprtf.striprtf import rtf_to_text
            
            # Decode RTF (usually encoded in ASCII/Latin-1)
            rtf_string = file_bytes.decode('latin-1', errors='ignore')
            
            # Strip RTF formatting
            plain_text = rtf_to_text(rtf_string)
            
            if not plain_text or not plain_text.strip():
                raise ValueError("RTF file appears to be empty")
            
            logger.info(f"Extracted {len(plain_text)} characters from .rtf")
            return plain_text.strip()
        
        except Exception as e:
            logger.error(f"Failed to extract text from .rtf: {e}")
            raise ValueError(f"Cannot parse .rtf file: {str(e)}")
    
    def _get_file_type(self, extension: str) -> FileType:
        """
        Determine file type from extension.
        
        Args:
            extension: File extension (with dot)
            
        Returns:
            FileType enum value
        """
        ext_lower = extension.lower()
        
        if ext_lower in self.TEXT_EXTENSIONS:
            return FileType.TEXT
        elif ext_lower in self.CONFIG_EXTENSIONS:
            return FileType.CONFIG
        elif ext_lower in self.SCRIPT_EXTENSIONS:
            return FileType.SCRIPT
        elif ext_lower in self.OFFICE_EXTENSIONS:
            return FileType.OFFICE
        elif ext_lower in self.RICH_EXTENSIONS:
            return FileType.RICH
        else:
            return FileType.UNKNOWN
    
    async def validate_file(
        self,
        filename: str,
        file_bytes: bytes
    ) -> Tuple[bool, Optional[str]]:
        """
        Validate uploaded file (size, type, content).
        
        Args:
            filename: Original filename
            file_bytes: Raw file bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check file size
            if len(file_bytes) == 0:
                return False, "File is empty"
            
            if len(file_bytes) > self.MAX_FILE_SIZE:
                size_mb = len(file_bytes) / (1024 * 1024)
                max_mb = self.MAX_FILE_SIZE / (1024 * 1024)
                return False, f"File too large ({size_mb:.2f}MB). Maximum: {max_mb}MB"
            
            # Check extension
            extension = Path(filename).suffix.lower()
            if not extension:
                return False, "File has no extension"
            
            file_type = self._get_file_type(extension)
            if file_type == FileType.UNKNOWN:
                return False, f"Unsupported file type: {extension}"
            
            # Check MIME type
            mime_type, _ = mimetypes.guess_type(filename)
            if mime_type and mime_type not in self.ALLOWED_MIME_TYPES:
                # Don't reject based on MIME alone, extension is more reliable
                logger.warning(f"Unrecognized MIME type {mime_type} for {filename}, but extension {extension} is supported")
            
            # For Office documents and rich formats, skip text validation (they're binary)
            if file_type in (FileType.OFFICE, FileType.RICH):
                logger.info(f"{file_type.value.capitalize()} file validation passed: {filename} ({len(file_bytes)} bytes, {extension})")
                return True, None
            
            # For text files, try to decode content
            encoding = self._detect_encoding(file_bytes)
            try:
                content = file_bytes.decode(encoding)
                if not content.strip():
                    return False, "File contains no readable text"
            except UnicodeDecodeError:
                return False, f"Cannot decode file with detected encoding: {encoding}"
            
            logger.info(f"File validation passed: {filename} ({len(file_bytes)} bytes, {extension}, {file_type.value})")
            return True, None
        
        except Exception as e:
            logger.error(f"File validation error: {e}")
            return False, f"Validation error: {str(e)}"
    
    async def parse_file(
        self,
        filename: str,
        file_bytes: bytes
    ) -> FileInfo:
        """
        Parse uploaded file and extract content.
        
        Args:
            filename: Original filename
            file_bytes: Raw file bytes
            
        Returns:
            FileInfo object with parsed data
            
        Raises:
            ValueError: If file is invalid
        """
        try:
            # Validate first
            is_valid, error_msg = await self.validate_file(filename, file_bytes)
            if not is_valid:
                raise ValueError(error_msg)
            
            # Extract file info
            extension = Path(filename).suffix.lower()
            file_type = self._get_file_type(extension)
            mime_type, _ = mimetypes.guess_type(filename)
            
            # Extract content based on file type
            if file_type == FileType.OFFICE:
                # Use specialized extractors for Office documents
                if extension == '.docx':
                    content = self._extract_text_from_docx(file_bytes)
                    encoding = 'utf-8'  # Extracted text is always UTF-8
                elif extension == '.xlsx':
                    content = self._extract_text_from_xlsx(file_bytes)
                    encoding = 'utf-8'
                elif extension == '.odt':
                    content = self._extract_text_from_odt(file_bytes)
                    encoding = 'utf-8'
                elif extension == '.ods':
                    content = self._extract_text_from_ods(file_bytes)
                    encoding = 'utf-8'
                elif extension in ('.doc', '.xls'):
                    # Old formats not yet supported
                    raise ValueError(f"Old Office format {extension} not yet supported. Please convert to .docx/.xlsx")
                else:
                    raise ValueError(f"Unsupported Office format: {extension}")
            elif file_type == FileType.RICH:
                # Use specialized extractors for rich formats
                if extension == '.pdf':
                    content = self._extract_text_from_pdf(file_bytes)
                    encoding = 'utf-8'  # Extracted text is always UTF-8
                elif extension == '.rtf':
                    content = self._extract_text_from_rtf(file_bytes)
                    encoding = 'utf-8'
                else:
                    raise ValueError(f"Unsupported rich format: {extension}")
            else:
                # Text-based files: decode with detected encoding
                encoding = self._detect_encoding(file_bytes)
                content = file_bytes.decode(encoding)
            
            file_info = FileInfo(
                filename=filename,
                size_bytes=len(file_bytes),
                mime_type=mime_type or 'text/plain',
                file_type=file_type,
                extension=extension,
                encoding=encoding,
                content=content
            )
            
            logger.info(f"File parsed successfully: {filename} ({file_type.value}, {encoding})")
            return file_info
        
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"File parsing error: {e}")
            raise ValueError(f"Failed to parse file: {str(e)}")
    
    async def export_file(
        self,
        original_filename: str,
        anonymized_content: str,
        encoding: str = 'utf-8'
    ) -> Tuple[str, bytes]:
        """
        Export anonymized content with original format.
        
        For binary formats (Office, PDF, RTF), converts to .txt since
        reconstructing the original format would lose anonymization mapping.
        
        Args:
            original_filename: Original file name
            anonymized_content: Anonymized text content
            encoding: File encoding (default: utf-8)
            
        Returns:
            Tuple of (new_filename, file_bytes)
        """
        try:
            path = Path(original_filename)
            stem = path.stem
            suffix = path.suffix.lower()
            
            # For binary formats (Office, PDF, RTF), export as .txt
            # Reconstructing the original format is complex and would require
            # maintaining formatting, which conflicts with text anonymization
            binary_formats = {'.docx', '.xlsx', '.odt', '.ods', '.pdf', '.rtf', '.doc', '.xls'}
            
            if suffix in binary_formats:
                # Convert to .txt for binary formats
                new_filename = f"{stem}.anonymized.txt"
                logger.info(f"Converting binary format {suffix} to .txt for export")
            else:
                # Keep original format for text-based files
                new_filename = f"{stem}.anonymized{suffix}"
            
            # Encode content
            file_bytes = anonymized_content.encode(encoding)
            
            logger.info(f"File exported: {new_filename} ({len(file_bytes)} bytes, {encoding})")
            return new_filename, file_bytes
        
        except Exception as e:
            logger.error(f"File export error: {e}")
            raise ValueError(f"Failed to export file: {str(e)}")
    
    def get_supported_extensions(self) -> Dict[str, list]:
        """
        Get list of all supported file extensions by category.
        
        Returns:
            Dictionary with categories as keys and extension lists as values
        """
        return {
            "text": sorted(list(self.TEXT_EXTENSIONS)),
            "config": sorted(list(self.CONFIG_EXTENSIONS)),
            "script": sorted(list(self.SCRIPT_EXTENSIONS)),
        }
    
    def get_file_size_limit(self) -> Dict[str, Any]:
        """
        Get file size limit information.
        
        Returns:
            Dictionary with size limits in different units
        """
        return {
            "bytes": self.MAX_FILE_SIZE,
            "kb": self.MAX_FILE_SIZE / 1024,
            "mb": self.MAX_FILE_SIZE / (1024 * 1024),
        }
